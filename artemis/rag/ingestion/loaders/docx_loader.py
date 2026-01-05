"""
DOCX loader for A.R.T.E.M.I.S.
"""

from pathlib import Path
from typing import Union

from artemis.utils import get_logger

logger = get_logger(__name__)


def load_docx_text(path: Union[str, Path]) -> str:
    """
    Load DOCX file and extract text content.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ImportError: If python-docx is not installed
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    
    logger.debug(f"Loading DOCX file: {path}")
    
    try:
        from docx import Document
        doc = Document(path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        text = "\n\n".join(text_parts)
        logger.info(f"Loaded DOCX with {len(doc.paragraphs)} paragraphs")
        return text
    except ImportError:
        raise ImportError(
            "DOCX loading requires 'python-docx'. "
            "Install with: pip install python-docx"
        )
    except Exception as e:
        logger.exception(f"Failed to load DOCX file: {path}", exc_info=True)
        raise

