"""
File loaders for A.R.T.E.M.I.S.

Provides file type-specific loaders that extract raw content from various
file formats. Handles missing dependencies gracefully.
"""

from pathlib import Path
from typing import Union

import pandas as pd

from artemis.utils import get_logger

logger = get_logger(__name__)


def load_csv(path: Union[str, Path]) -> pd.DataFrame:
    """
    Load CSV file as pandas DataFrame.
    
    Args:
        path: Path to CSV file
        
    Returns:
        pandas DataFrame
        
    Raises:
        FileNotFoundError: If file doesn't exist
        pd.errors.EmptyDataError: If CSV is empty
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"CSV file not found: {path}")
    
    logger.debug(f"Loading CSV file: {path}")
    try:
        df = pd.read_csv(path)
        logger.info(f"Loaded CSV with {len(df)} rows and {len(df.columns)} columns")
        return df
    except Exception as e:
        logger.exception(f"Failed to load CSV file: {path}", exc_info=True)
        raise


def load_pdf_text(path: Union[str, Path]) -> str:
    """
    Load PDF file and extract text content.
    
    Args:
        path: Path to PDF file
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ImportError: If required PDF library is not installed
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")
    
    logger.debug(f"Loading PDF file: {path}")
    
    # Try pdfplumber first (better text extraction)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts)
        logger.info(f"Loaded PDF with {len(pdf.pages)} pages using pdfplumber")
        return text
    except ImportError:
        logger.debug("pdfplumber not available, trying PyPDF2")
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        text_parts = []
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts)
        logger.info(f"Loaded PDF with {len(pdf_reader.pages)} pages using PyPDF2")
        return text
    except ImportError:
        raise ImportError(
            "PDF loading requires either 'pdfplumber' or 'PyPDF2'. "
            "Install with: pip install pdfplumber or pip install PyPDF2"
        )
    except Exception as e:
        logger.exception(f"Failed to load PDF file: {path}", exc_info=True)
        raise


def load_docx_text(path: Union[str, Path]) -> str:
    """
    Load DOCX file and extract text content.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ImportError: If python-docx is not installed
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    
    logger.debug(f"Loading DOCX file: {path}")
    
    try:
        from docx import Document
        doc = Document(path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        text = "\n\n".join(text_parts)
        logger.info(f"Loaded DOCX with {len(doc.paragraphs)} paragraphs")
        return text
    except ImportError:
        raise ImportError(
            "DOCX loading requires 'python-docx'. "
            "Install with: pip install python-docx"
        )
    except Exception as e:
        logger.exception(f"Failed to load DOCX file: {path}", exc_info=True)
        raise


def load_md_text(path: Union[str, Path]) -> str:
    """
    Load Markdown file and return text content.
    
    Args:
        path: Path to Markdown file
        
    Returns:
        File content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Markdown file not found: {path}")
    
    logger.debug(f"Loading Markdown file: {path}")
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        logger.info(f"Loaded Markdown file: {len(text)} characters")
        return text
    except Exception as e:
        logger.exception(f"Failed to load Markdown file: {path}", exc_info=True)
        raise


def load_text(path: Union[str, Path]) -> str:
    """
    Load plain text file.
    
    Args:
        path: Path to text file
        
    Returns:
        File content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Text file not found: {path}")
    
    logger.debug(f"Loading text file: {path}")
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        logger.info(f"Loaded text file: {len(text)} characters")
        return text
    except Exception as e:
        logger.exception(f"Failed to load text file: {path}", exc_info=True)
        raise

